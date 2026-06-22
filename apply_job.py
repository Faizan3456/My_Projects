"""
Single-job application bot.

Usage:
    python apply_job.py <linkedin_job_url>

What it does:
1. Opens browser with your saved LinkedIn session
2. Scrapes the job title, company, and description
3. Calls Claude to score the job and produce a tailored CV + cover letter
4. Saves a tailored CV DOCX (and PDF if Word is available)
5a. Easy Apply  → fully automates the form and submits
5b. External    → opens the apply URL, pauses for you to log in,
                  then auto-fills the form and waits for your final review before submitting
"""

import os, sys, re, time, random, json, datetime, argparse, subprocess
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

# ── dependencies ──────────────────────────────────────────────────────────────
try:
    import requests
except ImportError:
    sys.exit("Missing: pip install requests")

try:
    import pdfplumber
except ImportError:
    sys.exit("Missing: pip install pdfplumber")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("Missing: pip install python-docx")

try:
    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
except ImportError:
    sys.exit("Missing: pip install playwright && playwright install chromium")

# ── config ────────────────────────────────────────────────────────────────────
CLAUDE_API_KEY   = os.getenv("CLAUDE_API_KEY", "")
CV_PDF_PATH      = os.getenv("CV_PATH", "Faizan_Fayyaz_KYNDRYL.pdf")
MASTER_PROMPT    = Path("MasterPrompt.txt").read_text(encoding="utf-8") if Path("MasterPrompt.txt").exists() else ""
PROFILE_DIR      = os.getenv("LINKEDIN_PROFILE_DIR", "./linkedin_profile")
OUTPUT_DIR       = Path("tailored_cvs")
OUTPUT_DIR.mkdir(exist_ok=True)

# ── helpers ───────────────────────────────────────────────────────────────────
def _delay(lo=0.5, hi=1.5):
    time.sleep(random.uniform(lo, hi))

def _safe_inner_text(locator, timeout=5000) -> str:
    try:
        locator.first.wait_for(state="attached", timeout=timeout)
        return (locator.first.inner_text(timeout=timeout) or "").strip()
    except Exception:
        return ""

# ── CV text extraction ────────────────────────────────────────────────────────
def read_cv_pdf(path: str) -> str:
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text += (page.extract_text() or "") + "\n"
    return text[:8000]

# ── Claude call ───────────────────────────────────────────────────────────────
def ask_claude(prompt: str, max_tokens: int = 3000) -> str:
    if not CLAUDE_API_KEY:
        sys.exit("Set CLAUDE_API_KEY in .env")
    headers = {
        "Content-Type": "application/json",
        "x-api-key": CLAUDE_API_KEY,
        "anthropic-version": "2023-06-01",
    }
    data = {
        "model": "claude-sonnet-4-6",
        "max_tokens": max_tokens,
        "messages": [{"role": "user", "content": prompt}],
    }
    for attempt in range(3):
        try:
            r = requests.post("https://api.anthropic.com/v1/messages", json=data, headers=headers, timeout=90)
            r.raise_for_status()
            return r.json()["content"][0]["text"]
        except Exception as e:
            if attempt == 2:
                sys.exit(f"Claude API error: {e}")
            time.sleep(2 ** attempt)
    return ""

def extract_json_block(text: str) -> dict:
    text = text.strip()
    for prefix in ("```json", "```"):
        if text.startswith(prefix):
            text = text[len(prefix):]
    if text.endswith("```"):
        text = text[:-3]
    start = text.find("{")
    if start == -1:
        return {}
    depth = 0
    for i in range(start, len(text)):
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                try:
                    return json.loads(text[start:i+1])
                except Exception:
                    return {}
    return {}

# ── Claude analysis ───────────────────────────────────────────────────────────
def analyse_job(job_title: str, company: str, description: str) -> dict:
    cv_text = read_cv_pdf(CV_PDF_PATH)
    prompt = f"""{MASTER_PROMPT}

You are helping Faizan apply to this specific job. Produce a tailored application.

JOB TITLE: {job_title}
COMPANY: {company}
JOB DESCRIPTION:
{description}

CANDIDATE CV:
{cv_text}

Return ONLY valid JSON with this exact structure:
{{
  "decision": "APPLY" or "SKIP",
  "score": 0-100,
  "reason": "one sentence",
  "sponsorship_probability": "Low" | "Medium" | "High",
  "tailored_summary": "2-3 sentence professional summary rewritten for THIS job using only real CV facts",
  "key_achievements": [
    "rewritten achievement bullet tailored to this JD — real facts only, no invention"
  ],
  "skills_to_highlight": ["skill1", "skill2"],
  "cover_letter": "full tailored cover letter (4 paragraphs)",
  "work_authorization_answer": "...",
  "visa_sponsorship_answer": "...",
  "salary_expectation_answer": "...",
  "notice_period_answer": "...",
  "location_answer": "..."
}}"""

    print("  Asking Claude to analyse job and tailor CV...")
    raw = ask_claude(prompt, max_tokens=3000)
    data = extract_json_block(raw)
    if not data:
        print("  WARNING: Claude returned unexpected format, using raw text")
        data = {"decision": "APPLY", "score": 50, "reason": "Parse error", "tailored_summary": raw[:500]}
    return data

# ── CV generation ─────────────────────────────────────────────────────────────
def build_tailored_cv_docx(analysis: dict, job_title: str, company: str) -> Path:
    """Create a tailored CV DOCX from the original CV with Claude's improvements."""
    cv_text = read_cv_pdf(CV_PDF_PATH)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = section.bottom_margin   = Pt(36)
        section.left_margin   = section.right_margin    = Pt(54)

    def heading(text, size=14, bold=True, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def body(text, size=10.5):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        return p

    def rule():
        p = doc.add_paragraph("─" * 80)
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # ── Header ──
    heading("Faizan Fayyaz", size=18, color=(0, 51, 102))
    body("London, UK  |  fznfayyaz@gmail.com  |  linkedin.com/in/faizanfayyaz")
    rule()

    # ── Tailored Summary ──
    heading("Professional Summary", size=12, color=(0, 51, 102))
    summary = analysis.get("tailored_summary", "")
    if summary:
        body(summary)

    # ── Key Achievements ──
    achievements = analysis.get("key_achievements", [])
    if achievements:
        rule()
        heading("Key Achievements", size=12, color=(0, 51, 102))
        for ach in achievements:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(ach).strip())
            run.font.size = Pt(10.5)

    # ── Skills ──
    skills = analysis.get("skills_to_highlight", [])
    if skills:
        rule()
        heading("Core Skills", size=12, color=(0, 51, 102))
        body("  ·  ".join(skills))

    # ── Original CV body (rest of experience etc.) ──
    rule()
    heading("Career History & Experience", size=12, color=(0, 51, 102))
    # Strip the top of the CV (name/contact block) — keep from first job onward
    lines = cv_text.split("\n")
    in_body = False
    for line in lines:
        if not in_body:
            # Start after the first occurrence of a year (job dates)
            if re.search(r"\b(20\d{2}|19\d{2})\b", line):
                in_body = True
        if in_body and line.strip():
            body(line)

    # ── Cover Letter (separate page) ──
    cover = analysis.get("cover_letter", "")
    if cover:
        doc.add_page_break()
        heading(f"Cover Letter — {job_title} at {company}", size=12, color=(0, 51, 102))
        rule()
        for para in cover.split("\n\n"):
            if para.strip():
                body(para.strip())

    # Save
    safe_company = re.sub(r"[^\w]", "_", company)[:30]
    safe_title   = re.sub(r"[^\w]", "_", job_title)[:30]
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    docx_path = OUTPUT_DIR / f"CV_{safe_company}_{safe_title}_{ts}.docx"
    doc.save(str(docx_path))
    print(f"  ✓ Tailored CV saved: {docx_path}")

    # Try to convert to PDF (macOS: uses Word via docx2pdf)
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"  ✓ PDF saved:         {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"  ℹ PDF conversion skipped ({e}) — using DOCX")
        return docx_path

# ── LinkedIn scraping ─────────────────────────────────────────────────────────
def scrape_job_page(page, url: str) -> dict:
    """Return {title, company, description, is_easy_apply, external_apply_url}"""
    print(f"  Loading job page...")
    page.goto(url, wait_until="domcontentloaded", timeout=30000)
    try:
        page.wait_for_load_state("networkidle", timeout=8000)
    except Exception:
        pass
    _delay(1.5, 2.5)

    if "authwall" in page.url or "login" in page.url:
        sys.exit("Not logged in — run: python linkedin_test.py  (log in once, then rerun)")

    # Tab title is most reliable: "Cloud Engineer at Accenture | LinkedIn"
    title, company = "", ""
    try:
        tab = re.sub(r"\s*\|.*$", "", page.title()).strip()
        if " at " in tab:
            title, company = tab.split(" at ", 1)
            title = title.strip(); company = company.strip()
        elif tab:
            title = tab
    except Exception:
        pass
    if not title:
        try:
            for h1 in page.locator("h1").all():
                t = (h1.inner_text(timeout=2000) or "").strip()
                if t and len(t) > 3:
                    title = t; break
        except Exception:
            pass
    if not company:
        try:
            company = page.evaluate("""() => {
                for (const el of document.querySelectorAll('a[href*="/company/"], [class*="company"]')) {
                    const t = el.innerText?.trim();
                    if (t && t.length > 1 && t.length < 80) return t;
                }
                return '';
            }""")
        except Exception:
            pass
    title   = title   or "Unknown Role"
    company = company or "Unknown Company"

    # Expand description
    try:
        btn = page.locator("button:has-text('Show more')").first
        if btn.is_visible(timeout=2000):
            btn.click()
            _delay(0.5, 1)
    except Exception:
        pass

    description = ""
    for sel in ["#job-details", ".jobs-description__content", ".jobs-description-content__text", "article"]:
        description = _safe_inner_text(page.locator(sel), timeout=4000)
        if len(description) > 100:
            break
    if not description:
        description = page.inner_text("body", timeout=5000)[:5000]

    # Detect Easy Apply vs external
    is_easy_apply = False
    external_url  = ""
    try:
        ea_btn = page.locator("button.jobs-apply-button:has-text('Easy Apply'), button[aria-label*='Easy Apply']").first
        if ea_btn.is_visible(timeout=3000):
            is_easy_apply = True
    except Exception:
        pass

    if not is_easy_apply:
        try:
            apply_btn = page.locator("button.jobs-apply-button, a.jobs-apply-button").first
            if apply_btn.is_visible(timeout=3000):
                href = apply_btn.get_attribute("href") or ""
                if href.startswith("http"):
                    external_url = href
        except Exception:
            pass

    return {
        "title": title or "Unknown Role",
        "company": company or "Unknown Company",
        "description": description,
        "is_easy_apply": is_easy_apply,
        "external_apply_url": external_url,
        "job_url": url,
    }

# ── Easy Apply automation ─────────────────────────────────────────────────────
def run_easy_apply(page, job: dict, analysis: dict, cv_path: Path):
    print("\n  [Easy Apply] Starting...")

    # Click Easy Apply button
    try:
        btn = page.locator("button.jobs-apply-button:has-text('Easy Apply'), button[aria-label*='Easy Apply']").first
        btn.wait_for(timeout=8000)
        btn.click()
        _delay(1.5, 2.5)
    except Exception as e:
        print(f"  ✗ Could not open Easy Apply: {e}")
        return False

    return _handle_easy_apply_modal(page, analysis, cv_path)

def _handle_easy_apply_modal(page, analysis: dict, cv_path: Path) -> bool:
    for step in range(20):
        _delay(0.8, 1.5)

        # ── Submitted? ──
        if page.locator("h3:has-text('application was sent'), div:has-text('application was sent')").is_visible():
            print("  ✓ Application submitted!")
            return True

        # ── Upload resume ──
        try:
            upload = page.locator("input[type='file']").first
            if upload.is_visible(timeout=1000) and os.path.exists(str(cv_path)):
                upload.set_input_files(str(cv_path))
                print(f"    Uploaded: {cv_path.name}")
                _delay(1, 2)
        except Exception:
            pass

        # ── Fill text / select / radio fields ──
        _fill_form_fields(page, analysis)

        # ── Advance ──
        advanced = _click_next_or_submit(page)
        if advanced == "submitted":
            print("  ✓ Application submitted!")
            return True
        elif advanced == "stuck":
            print(f"  ✗ Stuck at step {step+1} — no button found")
            return False

    print("  ✗ Reached max steps without submitting")
    return False

def _fill_form_fields(page, analysis: dict):
    answers = {
        "phone":               "07700000000",
        "mobile":              "07700000000",
        "salary":              analysis.get("salary_expectation_answer", "£60,000"),
        "expected salary":     analysis.get("salary_expectation_answer", "£60,000"),
        "notice":              analysis.get("notice_period_answer", "1 month"),
        "availability":        analysis.get("notice_period_answer", "1 month"),
        "years of experience": "7",
        "how many years":      "7",
        "city":                "London",
        "linkedin":            "https://www.linkedin.com/in/faizanfayyaz",
        "cover letter":        analysis.get("cover_letter", "")[:1000],
        "additional":          analysis.get("cover_letter", "")[:500],
    }

    # Text inputs
    try:
        inputs = page.locator("input[type='text'], input[type='number'], input[type='tel'], textarea").all()
        for inp in inputs:
            try:
                aria  = (inp.get_attribute("aria-label") or "").lower()
                ph    = (inp.get_attribute("placeholder") or "").lower()
                label = aria or ph
                current = inp.input_value()
                if current.strip():
                    continue  # already filled
                for key, val in answers.items():
                    if key in label and val:
                        inp.click()
                        inp.fill(val)
                        _delay(0.1, 0.3)
                        break
            except Exception:
                pass
    except Exception:
        pass

    # Work auth / right to work → Yes radio
    for phrase in ["authorised to work", "right to work", "eligible to work"]:
        _click_radio_containing(page, phrase, "yes")

    # Sponsorship → Yes (we need sponsorship)
    for phrase in ["require.*sponsor", "need.*sponsor", "visa sponsor"]:
        _click_radio_containing(page, phrase, "yes")

    # Fill any still-empty required inputs with a safe default
    try:
        for inp in page.locator("input[required]:not([type='file']):not([type='radio']):not([type='checkbox']):not([type='hidden'])").all():
            try:
                if not inp.input_value().strip():
                    inp.fill("N/A")
            except Exception:
                pass
    except Exception:
        pass

    # Select first option for any untouched dropdowns
    try:
        for sel in page.locator("select").all():
            try:
                if not sel.input_value():
                    opts = sel.locator("option").all()
                    if len(opts) > 1:
                        sel.select_option(index=1)
            except Exception:
                pass
    except Exception:
        pass

def _click_radio_containing(page, phrase_regex: str, choice: str):
    try:
        fieldsets = page.locator("fieldset").all()
        for fs in fieldsets:
            legend = (fs.locator("legend").first.inner_text(timeout=500) or "").lower()
            if re.search(phrase_regex, legend):
                for radio in fs.locator("input[type='radio']").all():
                    rid = radio.get_attribute("id") or ""
                    lbl = (page.locator(f"label[for='{rid}']").first.inner_text(timeout=500) or "").lower()
                    if choice in lbl:
                        radio.check()
                        _delay(0.1, 0.3)
                        return
    except Exception:
        pass

def _click_next_or_submit(page) -> str:
    _delay(0.3, 0.7)
    for selector, label in [
        ("button[aria-label*='Submit application'], button:has-text('Submit application')", "submit"),
        ("button[aria-label*='Review your application'], button:has-text('Review')", "review"),
        ("button[aria-label*='Continue to next step'], button:has-text('Next')", "next"),
        ("button:has-text('Done')", "done"),
    ]:
        try:
            btn = page.locator(selector).first
            if btn.is_visible(timeout=1500):
                btn.click()
                _delay(1, 2)
                if label == "submit":
                    return "submitted"
                return "next"
        except Exception:
            pass

    # Dismiss follow/save popups
    try:
        dismiss = page.locator("button:has-text('Not now'), button:has-text('Skip'), button:has-text('Dismiss')").first
        if dismiss.is_visible(timeout=1000):
            dismiss.click()
            return "next"
    except Exception:
        pass

    return "stuck"

# ── External / third-party apply ──────────────────────────────────────────────
def run_external_apply(page, job: dict, analysis: dict, cv_path: Path):
    print("\n  [External ATS] Opening application page...")

    # Click the Apply button to get the external URL (it may open a new tab)
    ext_url = job.get("external_apply_url", "")

    # Go back to job page first
    page.goto(job["job_url"], wait_until="domcontentloaded", timeout=20000)
    _delay(1.5, 2)

    new_page = None
    try:
        with page.context.expect_page(timeout=8000) as page_info:
            apply_btn = page.locator("button.jobs-apply-button, a.jobs-apply-button").first
            apply_btn.click()
        new_page = page_info.value
        new_page.wait_for_load_state("domcontentloaded", timeout=15000)
        ext_url = new_page.url
    except Exception:
        if ext_url:
            new_page = page.context.new_page()
            new_page.goto(ext_url, wait_until="domcontentloaded", timeout=20000)
        else:
            print("  ✗ Could not find external apply URL")
            return False

    target = new_page or page
    print(f"  External URL: {target.url}")

    # ── Pause for manual login if needed ──
    if _needs_login(target):
        print("\n" + "="*60)
        print("  ACTION REQUIRED: Please log in on the browser window.")
        print("  Once you're logged in and the application page is open,")
        print("  press Enter here to continue...")
        print("="*60)
        input()
        _delay(1, 2)

    print("  Filling in the application form...")
    _fill_form_fields(target, analysis)

    # Try to upload CV
    try:
        upload = target.locator("input[type='file']").first
        if upload.is_visible(timeout=3000) and os.path.exists(str(cv_path)):
            upload.set_input_files(str(cv_path))
            print(f"  Uploaded CV: {cv_path.name}")
            _delay(1, 2)
    except Exception:
        pass

    # Fill cover letter textarea if present
    cover = analysis.get("cover_letter", "")
    if cover:
        for sel in ["textarea[name*='cover'], textarea[id*='cover'], textarea[placeholder*='cover']",
                    "textarea"]:
            try:
                ta = target.locator(sel).first
                if ta.is_visible(timeout=2000):
                    if not ta.input_value().strip():
                        ta.fill(cover[:2000])
                        _delay(0.3, 0.7)
                        break
            except Exception:
                pass

    print("\n" + "="*60)
    print("  REVIEW: The form has been filled. Please check it in the browser.")
    print("  Fix anything that looks wrong, then press Enter here to submit.")
    print("="*60)
    input()

    # Click submit
    for sel in ["button[type='submit'], input[type='submit'], button:has-text('Submit'), button:has-text('Apply')"]:
        try:
            btn = target.locator(sel).first
            if btn.is_visible(timeout=3000):
                btn.click()
                print("  ✓ Submitted!")
                _delay(2, 3)
                return True
        except Exception:
            pass

    print("  Could not find submit button — please click it manually.")
    input("  Press Enter after submitting...")
    return True

def _needs_login(page) -> bool:
    url = page.url.lower()
    login_signals = ["login", "signin", "sign-in", "auth", "sso", "saml"]
    if any(s in url for s in login_signals):
        return True
    try:
        page.locator("input[type='password']").first.wait_for(state="visible", timeout=3000)
        return True
    except Exception:
        return False

# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="Apply to a single LinkedIn job")
    parser.add_argument("url", help="LinkedIn job URL (e.g. https://www.linkedin.com/jobs/view/...)")
    parser.add_argument("--skip-cv", action="store_true", help="Skip CV tailoring, use original PDF")
    parser.add_argument("--dry-run", action="store_true", help="Analyse and generate CV but don't apply")
    args = parser.parse_args()

    print(f"\n{'='*60}")
    print(f"  Job Application Bot")
    print(f"  URL: {args.url}")
    print(f"{'='*60}\n")

    with sync_playwright() as pw:
        browser = pw.chromium.launch_persistent_context(
            user_data_dir=PROFILE_DIR,
            headless=False,
            args=["--disable-blink-features=AutomationControlled"],
            viewport={"width": 1280, "height": 900},
        )
        page = browser.new_page()

        # ── Step 1: Scrape job ──
        print("[1/4] Scraping job details...")
        job = scrape_job_page(page, args.url)
        print(f"  Title:   {job['title']}")
        print(f"  Company: {job['company']}")
        print(f"  Type:    {'Easy Apply' if job['is_easy_apply'] else 'External ATS'}")
        print(f"  Desc:    {len(job['description'])} chars")

        if not job["description"] or len(job["description"]) < 50:
            print("  ✗ Could not scrape job description. Aborting.")
            browser.close()
            return

        # ── Step 2: Claude analysis ──
        print("\n[2/4] Analysing job with Claude...")
        analysis = analyse_job(job["title"], job["company"], job["description"])
        print(f"  Decision:       {analysis.get('decision', '?')}")
        print(f"  Score:          {analysis.get('score', '?')}/100")
        print(f"  Sponsorship:    {analysis.get('sponsorship_probability', '?')}")
        print(f"  Reason:         {analysis.get('reason', '')}")

        if analysis.get("decision", "APPLY").upper() == "SKIP":
            print(f"\n  Claude recommends SKIP. Exiting.")
            browser.close()
            return

        # ── Step 3: Build tailored CV ──
        if args.skip_cv:
            cv_path = Path(CV_PDF_PATH)
            print(f"\n[3/4] Using original CV: {cv_path}")
        else:
            print("\n[3/4] Building tailored CV...")
            cv_path = build_tailored_cv_docx(analysis, job["title"], job["company"])

        if args.dry_run:
            print("\n  --dry-run: stopping before application. CV saved above.")
            browser.close()
            return

        # ── Step 4: Apply ──
        print(f"\n[4/4] Applying...")

        # Navigate back to job page (scraping may have moved us)
        page.goto(args.url, wait_until="domcontentloaded", timeout=20000)
        _delay(1.5, 2.5)

        if job["is_easy_apply"]:
            success = run_easy_apply(page, job, analysis, cv_path)
        else:
            success = run_external_apply(page, job, analysis, cv_path)

        status = "SUBMITTED" if success else "FAILED"
        print(f"\n{'='*60}")
        print(f"  Result: {status}")
        print(f"  CV used: {cv_path}")
        print(f"{'='*60}")

        # Log to applications.csv
        import csv
        log_path = "applications.csv"
        file_exists = Path(log_path).exists()
        with open(log_path, "a", newline="", encoding="utf-8") as f:
            w = csv.DictWriter(f, fieldnames=["timestamp","company","title","url","score","sponsorship","status","cv_path"])
            if not file_exists:
                w.writeheader()
            w.writerow({
                "timestamp":    datetime.datetime.utcnow().isoformat(),
                "company":      job["company"],
                "title":        job["title"],
                "url":          args.url,
                "score":        analysis.get("score", ""),
                "sponsorship":  analysis.get("sponsorship_probability", ""),
                "status":       status,
                "cv_path":      str(cv_path),
            })

        input("\nPress Enter to close the browser...")
        browser.close()

if __name__ == "__main__":
    main()
