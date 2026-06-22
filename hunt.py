"""
hunt.py — Full autonomous job hunt: Claude generates searches from your CV +
MasterPrompt → LinkedIn scrape → Claude scores all jobs → applies to top N.

Usage:
    python hunt.py                    # full run: search → score → apply top 20
    python hunt.py --score-only       # search + score, show ranking, no apply
    python hunt.py --top 10           # apply top 10 instead of 20
    python hunt.py --min-score 65     # override score threshold
    python hunt.py --limit 25         # collect more jobs per search term
"""

import os, sys, re, time, random, json, csv, datetime, argparse, concurrent.futures
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

try:
    import requests
except ImportError:
    sys.exit("pip install requests")
try:
    import pdfplumber
except ImportError:
    sys.exit("pip install pdfplumber")
try:
    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
except ImportError:
    sys.exit("pip install playwright && playwright install chromium")

# ── config ────────────────────────────────────────────────────────────────────
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY", "")
CV_PDF_PATH    = os.getenv("CV_PATH", "Faizan_Fayyaz_KYNDRYL.pdf")
MASTER_PROMPT  = Path("MasterPrompt.txt").read_text(encoding="utf-8") if Path("MasterPrompt.txt").exists() else ""
PROFILE_DIR    = os.getenv("LINKEDIN_PROFILE_DIR", "./linkedin_profile")
RESULTS_DIR    = Path("hunt_results"); RESULTS_DIR.mkdir(exist_ok=True)

LOCATION = "United Kingdom"

# Fallback searches used only if Claude search-generation fails
FALLBACK_SEARCHES = [
    "Application Support Engineer banking",
    "Production Support Engineer payments",
    "Technical Support Engineer SWIFT",
    "Infrastructure Support Engineer",
    "Trading Support Engineer",
    "Settlements Support Engineer",
    "IBM MQ Support Engineer",
    "Core Banking Support",
    "Payments Support Engineer",
    "Banking Technology Support",
]

# ── helpers ───────────────────────────────────────────────────────────────────
def _delay(lo=0.3, hi=0.8): time.sleep(random.uniform(lo, hi))
def _long_delay(lo=1.5, hi=3.0): time.sleep(random.uniform(lo, hi))  # only between applications

# Titles containing any of these words are skipped without visiting the page
_TITLE_BLOCKLIST = re.compile(
    r"\b(sales|marketing|product manager|data scientist|front.?end|react|angular|"
    r"software engineer|developer|devops engineer|desktop support|field service|"
    r"customer success|account manager|business development|recruiter|"
    r"hr |human resources|finance manager|scrum master|project manager|"
    r"game|graphic|design|seo|content|copywriter|pharmacist|nurse|teacher)\b",
    re.IGNORECASE
)

def _title_passes(title: str) -> bool:
    """Return False if the title is obviously out of scope."""
    return not bool(_TITLE_BLOCKLIST.search(title))

def _safe_text(locator, timeout=5000) -> str:
    try:
        locator.first.wait_for(state="attached", timeout=timeout)
        return (locator.first.inner_text(timeout=timeout) or "").strip()
    except Exception:
        return ""

def read_cv() -> str:
    text = ""
    with pdfplumber.open(CV_PDF_PATH) as pdf:
        for p in pdf.pages:
            text += (p.extract_text() or "") + "\n"
    return text[:7000]

# ── Smart search generation ───────────────────────────────────────────────────
def generate_searches(cv_text: str, n: int = 20) -> list:
    """
    Ask Claude to read the MasterPrompt + CV and produce the best LinkedIn
    search queries to find matching roles. Returns a list of search strings.
    """
    prompt = f"""{MASTER_PROMPT}

You are generating LinkedIn job search queries for Faizan.

Read his CV and the role preferences above, then produce exactly {n} LinkedIn
search keyword strings that will find the most relevant job postings for him.

Rules:
- Each query should be 2-5 words that a recruiter would use in a job title or skill tag
- Cover the range of his experience: application support, production support, payments, SWIFT, infrastructure, banking technology, IBM MQ, ServiceNow, monitoring
- Vary the queries — don't repeat the same role with slightly different words
- Target UK banking, payments, fintech, and enterprise IT sectors
- Include both role-title queries ("Application Support Engineer") and technology queries ("SWIFT Messaging Support")
- Do NOT include location — that is added separately

Return ONLY a JSON array of strings, no other text:
["query 1", "query 2", ...]

CANDIDATE CV:
{cv_text[:3000]}"""

    raw = ask_claude(prompt, max_tokens=800)
    raw = raw.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
    try:
        queries = json.loads(raw)
        if isinstance(queries, list) and all(isinstance(q, str) for q in queries):
            return [q.strip() for q in queries if q.strip()][:n]
    except Exception:
        pass
    # Try extracting array from mixed text
    m = re.search(r"\[.*?\]", raw, re.DOTALL)
    if m:
        try:
            queries = json.loads(m.group())
            if isinstance(queries, list):
                return [str(q).strip() for q in queries if str(q).strip()][:n]
        except Exception:
            pass
    return []

# ── Claude ────────────────────────────────────────────────────────────────────
def ask_claude(prompt: str, max_tokens: int = 600) -> str:
    headers = {"Content-Type": "application/json", "x-api-key": CLAUDE_API_KEY, "anthropic-version": "2023-06-01"}
    data = {"model": "claude-haiku-4-5-20251001", "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}]}
    for attempt in range(3):
        try:
            r = requests.post("https://api.anthropic.com/v1/messages", json=data, headers=headers, timeout=60)
            r.raise_for_status()
            return r.json()["content"][0]["text"]
        except Exception:
            time.sleep(2 ** attempt)
    return ""

def ask_claude_full(prompt: str, max_tokens: int = 3000) -> str:
    """Use Sonnet for full application package generation."""
    headers = {"Content-Type": "application/json", "x-api-key": CLAUDE_API_KEY, "anthropic-version": "2023-06-01"}
    data = {"model": "claude-sonnet-4-6", "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}]}
    for attempt in range(3):
        try:
            r = requests.post("https://api.anthropic.com/v1/messages", json=data, headers=headers, timeout=90)
            r.raise_for_status()
            return r.json()["content"][0]["text"]
        except Exception:
            time.sleep(2 ** attempt)
    return ""

def parse_json(text: str) -> dict:
    text = text.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
    start = text.find("{")
    if start == -1: return {}
    depth = 0
    for i in range(start, len(text)):
        if text[i] == "{": depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                try: return json.loads(text[start:i+1])
                except: return {}
    return {}

# ── Score a single job (fast, uses Haiku) ────────────────────────────────────
def score_job(job: dict, cv_text: str) -> dict:
    prompt = f"""{MASTER_PROMPT}

Score this job for Faizan using the scoring formula above. Return ONLY JSON.

IMPORTANT SCORING RULES:
- The posting company may be a recruitment agency (Lorien, Anson McCade, Hays, etc.) — this does NOT reduce the score. Score the actual role, not the recruiter.
- If salary is not mentioned, assume it may meet the minimum — do NOT penalise for missing salary info.
- If sponsorship is not mentioned but the employer is a large enterprise or bank, assume Medium sponsorship probability (50).
- "Contract" roles where right-to-work is explicitly required = SKIP. Otherwise score normally.
- Focus on: does the role match Faizan's banking/payments/infrastructure support background?

JOB TITLE: {job['title']}
POSTING COMPANY (may be recruiter): {job['company']}
JOB DESCRIPTION:
{job['description'][:3500]}

CANDIDATE CV (key facts):
{cv_text[:2000]}

Return ONLY this JSON (no other text):
{{
  "decision": "APPLY" or "SKIP",
  "score": 0-100,
  "role_alignment": 0-100,
  "technology_alignment": 0-100,
  "sponsorship_probability": 0-100,
  "salary_probability": 0-100,
  "company_quality": 0-100,
  "reason": "one sentence max"
}}"""

    raw = ask_claude(prompt, max_tokens=400)
    data = parse_json(raw)
    job["score"]                  = int(data.get("score", 0))
    job["decision"]               = str(data.get("decision", "SKIP")).upper()
    job["role_alignment"]         = data.get("role_alignment", 0)
    job["technology_alignment"]   = data.get("technology_alignment", 0)
    job["sponsorship_probability"]= data.get("sponsorship_probability", 0)
    job["salary_probability"]     = data.get("salary_probability", 0)
    job["company_quality"]        = data.get("company_quality", 0)
    job["reason"]                 = data.get("reason", "")
    return job

# ── Full application package (Sonnet, only for approved jobs) ─────────────────
def build_application_package(job: dict, cv_text: str) -> dict:
    prompt = f"""{MASTER_PROMPT}

Produce a full tailored application package for Faizan for this specific job.
Use ONLY real facts from his CV. Never invent anything.

JOB TITLE: {job['title']}
COMPANY: {job['company']}
JOB DESCRIPTION:
{job['description'][:4000]}

FULL CV:
{cv_text}

Return ONLY valid JSON:
{{
  "tailored_summary": "2-3 sentence summary rewritten for this role",
  "key_achievements": ["bullet 1", "bullet 2", "bullet 3", "bullet 4", "bullet 5"],
  "skills_to_highlight": ["skill1", "skill2", "skill3"],
  "cover_letter": "full 4-paragraph cover letter",
  "work_authorization_answer": "...",
  "visa_sponsorship_answer": "...",
  "salary_expectation_answer": "...",
  "notice_period_answer": "...",
  "location_answer": "..."
}}"""

    raw = ask_claude_full(prompt, max_tokens=2500)
    return parse_json(raw)

# ── LinkedIn scraping ─────────────────────────────────────────────────────────
def build_search_url(keywords: str) -> str:
    import urllib.parse
    kw  = urllib.parse.quote_plus(keywords)
    loc = urllib.parse.quote_plus(LOCATION)
    # Easy Apply + posted in last 7 days
    return f"https://www.linkedin.com/jobs/search/?keywords={kw}&location={loc}&f_TPR=r604800&f_LF=f_AL"

def collect_job_links(page, search_url: str, limit: int = 25) -> list:
    """Return unique /jobs/view/ URLs from a search results page."""
    try:
        page.goto(search_url, wait_until="domcontentloaded", timeout=20000)
        _delay(1.0, 1.5)   # no networkidle — just let JS settle briefly
    except Exception as e:
        print(f"    Search page error: {e}")
        return []

    if "authwall" in page.url or "login" in page.url:
        sys.exit("Not logged in — run: python linkedin_test.py")

    seen, urls = set(), []
    for _ in range(3):
        # Extract all job view hrefs via JS in one shot (faster than iterating locators)
        hrefs = page.evaluate("""() => {
            return Array.from(document.querySelectorAll('a[href*="/jobs/view/"]'))
                        .map(a => a.href);
        }""")
        for href in (hrefs or []):
            m = re.search(r"(https://www\.linkedin\.com/jobs/view/\d+)", href)
            if m:
                clean = m.group(1)
                if clean not in seen:
                    seen.add(clean); urls.append(clean)
            if len(urls) >= limit:
                return urls
        page.evaluate("window.scrollBy(0, document.body.scrollHeight)")
        _delay(1.0, 1.5)
        try:
            btn = page.locator("button:has-text('See more jobs')").first
            if btn.is_visible(timeout=500): btn.click(); _delay(0.8, 1.2)
        except Exception:
            pass
    return urls

def _extract_title_company(page) -> tuple:
    """
    Extract job title and company from a LinkedIn job page.
    Strategy 1: parse the browser tab title  e.g. "Cloud Engineer at Accenture | LinkedIn"
    Strategy 2: first <h1> on the page
    Strategy 3: aria-label on the top card
    """
    title, company = "", ""

    # Strategy 1 — browser tab title (most reliable, doesn't depend on DOM classes)
    try:
        tab_title = page.title()                     # "Cloud Engineer at Accenture | LinkedIn"
        tab_title = re.sub(r"\s*\|.*$", "", tab_title).strip()   # "Cloud Engineer at Accenture"
        if " at " in tab_title:
            parts = tab_title.split(" at ", 1)
            title   = parts[0].strip()
            company = parts[1].strip()
        elif tab_title:
            title = tab_title
    except Exception:
        pass

    # Strategy 2 — first visible h1
    if not title:
        try:
            h1s = page.locator("h1").all()
            for h1 in h1s:
                t = (h1.inner_text(timeout=2000) or "").strip()
                if t and len(t) > 3:
                    title = t; break
        except Exception:
            pass

    # Strategy 3 — JS fallback: find any element whose text contains "at <Company>"
    if not company:
        try:
            company = page.evaluate("""() => {
                const candidates = document.querySelectorAll('a[href*="/company/"], [class*="company"]');
                for (const el of candidates) {
                    const t = el.innerText?.trim();
                    if (t && t.length > 1 && t.length < 80) return t;
                }
                return '';
            }""")
        except Exception:
            pass

    return (title or "Unknown Role", company or "Unknown Company")


def _is_challenge_page(page) -> bool:
    """Return True if LinkedIn is showing a bot-check or auth-wall."""
    url = page.url.lower()
    return any(s in url for s in ["checkpoint", "authwall", "challenge", "/login", "/uas/"])

def scrape_job_page(page, url: str):
    """Scrape one job page. Returns dict or None on failure."""
    for attempt in range(2):
        try:
            page.goto(url, wait_until="domcontentloaded", timeout=25000)
            _delay(1.2, 2.0)   # human-like pause — prevents LinkedIn rate-limit
        except Exception:
            return None

        if _is_challenge_page(page):
            if attempt == 0:
                print(f"\n  ⚠ LinkedIn challenge detected — waiting 15s...", end="", flush=True)
                time.sleep(15)
                continue
            return None
        break

    # Click "Show more" if present
    try:
        btn = page.locator("button:has-text('Show more')").first
        if btn.is_visible(timeout=800): btn.click(); _delay(0.3, 0.5)
    except Exception:
        pass

    # Extract everything in one JS call
    data = page.evaluate("""() => {
        // Title + company from tab title (most reliable)
        let title = '', company = '';
        const tab = document.title.replace(/\\s*\\|.*$/, '').trim();
        if (tab.includes(' at ')) {
            const parts = tab.split(' at ');
            title   = parts[0].trim();
            company = parts.slice(1).join(' at ').trim();
        } else {
            title = tab;
        }
        // Fallback title: first h1
        if (!title) {
            const h1 = document.querySelector('h1');
            if (h1) title = h1.innerText.trim();
        }
        // Description: try IDs / class patterns in priority order
        const descSelectors = [
            '#job-details',
            '.jobs-description__content',
            '.jobs-description-content__text',
            '[class*="description__text"]',
            'article',
        ];
        let description = '';
        for (const sel of descSelectors) {
            const el = document.querySelector(sel);
            if (el && el.innerText.trim().length > 150) {
                description = el.innerText.trim().slice(0, 5000);
                break;
            }
        }
        // Easy Apply detection
        const buttons = Array.from(document.querySelectorAll('button'));
        const isEasyApply = buttons.some(b =>
            b.innerText.includes('Easy Apply') && b.offsetParent !== null
        );
        return { title, company, description, isEasyApply };
    }""")

    if not data or len(data.get("description", "")) < 100:
        return None

    return {
        "url":          url,
        "title":        data.get("title")       or "Unknown Role",
        "company":      data.get("company")     or "Unknown Company",
        "description":  data.get("description") or "",
        "is_easy_apply": bool(data.get("isEasyApply")),
    }

# ── Easy Apply ────────────────────────────────────────────────────────────────
def easy_apply(page, job: dict, pkg: dict, cv_path: str):
    print(f"      Clicking Easy Apply...")
    try:
        btn = page.locator("button.jobs-apply-button:has-text('Easy Apply'), button[aria-label*='Easy Apply']").first
        btn.wait_for(timeout=6000); btn.click(); _delay(1.5, 2.5)
    except Exception as e:
        print(f"      ✗ {e}"); return False

    for step in range(20):
        _delay(0.8, 1.5)

        if page.locator("h3:has-text('application was sent'), div:has-text('application was sent')").is_visible():
            return True

        # Upload CV
        try:
            f = page.locator("input[type='file']").first
            if f.is_visible(timeout=800) and os.path.exists(cv_path):
                f.set_input_files(cv_path); _delay(1, 1.5)
        except Exception:
            pass

        _fill_form(page, pkg)

        result = _advance(page)
        if result == "submitted": return True
        if result == "stuck":
            print(f"      Stuck at step {step+1}"); return False

    return False

def _fill_form(page, pkg: dict):
    answers = {
        "phone": "07700000000", "mobile": "07700000000",
        "salary": pkg.get("salary_expectation_answer", "£60,000"),
        "expected salary": pkg.get("salary_expectation_answer", "£60,000"),
        "notice": pkg.get("notice_period_answer", "1 month"),
        "availability": pkg.get("notice_period_answer", "1 month"),
        "years of experience": "7", "how many years": "7",
        "city": "London",
        "linkedin": "https://www.linkedin.com/in/faizanfayyaz",
        "cover letter": pkg.get("cover_letter", "")[:1000],
        "additional": pkg.get("cover_letter", "")[:500],
    }
    try:
        for inp in page.locator("input[type='text'],input[type='number'],input[type='tel'],textarea").all():
            try:
                label = ((inp.get_attribute("aria-label") or "") + " " + (inp.get_attribute("placeholder") or "")).lower()
                if inp.input_value().strip(): continue
                for k, v in answers.items():
                    if k in label and v:
                        inp.click(); inp.fill(v); _delay(0.1, 0.2); break
            except Exception:
                pass
    except Exception:
        pass
    # Radios: work auth → Yes
    for phrase in ["authorised to work", "right to work", "eligible to work"]:
        _radio(page, phrase, "yes")
    # Sponsorship → Yes
    for phrase in ["require.*sponsor", "need.*sponsor", "visa sponsor"]:
        _radio(page, phrase, "yes")
    # Empty required fields
    try:
        for inp in page.locator("input[required]:not([type='file']):not([type='radio']):not([type='hidden'])").all():
            try:
                if not inp.input_value().strip(): inp.fill("N/A")
            except Exception:
                pass
    except Exception:
        pass
    # Untouched selects
    try:
        for sel in page.locator("select").all():
            try:
                if not sel.input_value():
                    opts = sel.locator("option").all()
                    if len(opts) > 1: sel.select_option(index=1)
            except Exception:
                pass
    except Exception:
        pass

def _radio(page, phrase_re: str, choice: str):
    try:
        for fs in page.locator("fieldset").all():
            legend = (fs.locator("legend").first.inner_text(timeout=400) or "").lower()
            if re.search(phrase_re, legend):
                for r in fs.locator("input[type='radio']").all():
                    rid = r.get_attribute("id") or ""
                    lbl = (page.locator(f"label[for='{rid}']").first.inner_text(timeout=400) or "").lower()
                    if choice in lbl:
                        r.check(); _delay(0.1, 0.2); return
    except Exception:
        pass

def _advance(page) -> str:
    for sel, tag in [
        ("button[aria-label*='Submit application'],button:has-text('Submit application')", "submitted"),
        ("button:has-text('Review'),button[aria-label*='Review']", "next"),
        ("button:has-text('Next'),button[aria-label*='Continue']", "next"),
        ("button:has-text('Done')", "next"),
    ]:
        try:
            b = page.locator(sel).first
            if b.is_visible(timeout=1200):
                b.click(); _delay(1, 1.5)
                return tag
        except Exception:
            pass
    try:
        d = page.locator("button:has-text('Not now'),button:has-text('Skip')").first
        if d.is_visible(timeout=800):
            d.click(); return "next"
    except Exception:
        pass
    return "stuck"

# ── External ATS handler ─────────────────────────────────────────────────────
# ATS-specific submit button selectors (tried in order)
ATS_SUBMIT_SELECTORS = [
    # Generic
    "button[type='submit']",
    "input[type='submit']",
    "button:has-text('Submit')",
    "button:has-text('Apply')",
    "button:has-text('Submit Application')",
    "button:has-text('Complete Application')",
    # Workday
    "[data-automation-id='bottom-navigation-next-button']",
    "[data-automation-id='pageFooterNextButton']",
    # Greenhouse
    "#submit_app",
    "input[value='Submit Application']",
    # Lever
    "button.postings-btn-submit",
    # SmartRecruiters
    "button[data-test='btn-job-application-submit']",
    # iCIMS
    "input[value='Submit']",
    # Taleo
    "a[title='Submit']",
]

ATS_NEXT_SELECTORS = [
    "button:has-text('Next')",
    "button:has-text('Continue')",
    "button:has-text('Save and Continue')",
    "[data-automation-id='bottom-navigation-next-button']",
    "[data-automation-id='pageFooterNextButton']",
    "a:has-text('Next')",
    "button[aria-label*='Next']",
]

def _detect_ats(url: str) -> str:
    url = url.lower()
    if "myworkdayjobs" in url or "workday" in url: return "Workday"
    if "greenhouse.io" in url or "boards.greenhouse" in url: return "Greenhouse"
    if "lever.co" in url or "jobs.lever" in url: return "Lever"
    if "smartrecruiters" in url: return "SmartRecruiters"
    if "icims" in url: return "iCIMS"
    if "successfactors" in url or "sapsf" in url: return "SuccessFactors"
    if "taleo" in url: return "Taleo"
    if "oracle" in url and "recruit" in url: return "Oracle"
    return "Unknown ATS"

def _needs_login(page) -> bool:
    url = page.url.lower()
    if any(s in url for s in ["login", "signin", "sign-in", "sso", "saml", "auth"]):
        return True
    try:
        page.locator("input[type='password']").first.wait_for(state="visible", timeout=2000)
        return True
    except Exception:
        return False

def _fill_ats_form(page, pkg: dict, cv_path: str):
    """Fill common ATS form fields with answers from the application package."""
    answers = {
        "first name":          "Faizan",
        "last name":           "Fayyaz",
        "full name":           "Faizan Fayyaz",
        "name":                "Faizan Fayyaz",
        "email":               "fznfayyaz@gmail.com",
        "phone":               "07700000000",
        "mobile":              "07700000000",
        "linkedin":            "https://www.linkedin.com/in/faizanfayyaz",
        "website":             "",
        "portfolio":           "",
        "city":                "London",
        "location":            "London, UK",
        "address":             "London, UK",
        "postcode":            "",
        "salary":              pkg.get("salary_expectation_answer", "£60,000 - £70,000"),
        "expected salary":     pkg.get("salary_expectation_answer", "£60,000 - £70,000"),
        "desired salary":      pkg.get("salary_expectation_answer", "£60,000 - £70,000"),
        "notice":              pkg.get("notice_period_answer", "1 month"),
        "availability":        pkg.get("notice_period_answer", "1 month"),
        "start date":          "1 month notice",
        "years of experience": "7",
        "how many years":      "7",
        "cover letter":        pkg.get("cover_letter", "")[:2000],
        "additional":          pkg.get("cover_letter", "")[:800],
        "summary":             pkg.get("tailored_summary", ""),
        "current employer":    "Kyndryl",
        "current company":     "Kyndryl",
        "current role":        "Senior Application Support Engineer",
        "current title":       "Senior Application Support Engineer",
    }

    # Text inputs and textareas
    try:
        for inp in page.locator("input[type='text'], input[type='email'], input[type='tel'], input[type='number'], textarea").all():
            try:
                aria  = (inp.get_attribute("aria-label") or "").lower()
                ph    = (inp.get_attribute("placeholder") or "").lower()
                name  = (inp.get_attribute("name") or "").lower()
                label_text = aria or ph or name
                current = inp.input_value()
                if current.strip():
                    continue   # already filled, don't overwrite
                for key, val in answers.items():
                    if key in label_text and val:
                        inp.click(); inp.fill(val); _delay(0.1, 0.3); break
            except Exception:
                continue
    except Exception:
        pass

    # Work auth → Yes
    for phrase in ["authoris", "right to work", "eligible to work", "work in the uk"]:
        _radio(page, phrase, "yes")

    # Sponsorship → Yes
    for phrase in ["require.*sponsor", "need.*sponsor", "visa sponsor", "sponsorship"]:
        _radio(page, phrase, "yes")

    # Gender / diversity — prefer "prefer not to say"
    for phrase in ["gender", "ethnicity", "disability", "veteran"]:
        try:
            for fs in page.locator("fieldset, select").all():
                try:
                    legend = (fs.locator("legend, label").first.inner_text(timeout=300) or "").lower()
                    if phrase in legend:
                        # Try to find "prefer not to say" option
                        for opt in fs.locator("option, input[type='radio']").all():
                            txt = (opt.inner_text(timeout=300) or opt.get_attribute("value") or "").lower()
                            if "prefer" in txt or "decline" in txt or "not to say" in txt:
                                if opt.tag_name() == "option":
                                    fs.select_option(value=opt.get_attribute("value"))
                                else:
                                    opt.check()
                                break
                except Exception:
                    continue
        except Exception:
            pass

    # Upload CV
    try:
        upload = page.locator("input[type='file']").first
        if upload.is_visible(timeout=2000) and os.path.exists(cv_path):
            upload.set_input_files(cv_path)
            print(f"        Uploaded: {Path(cv_path).name}")
            _delay(1, 2)
    except Exception:
        pass

def _click_ats_next(page) -> str:
    """Click Next/Continue on multi-step ATS forms. Returns 'next', 'submit', or 'stuck'."""
    _delay(0.5, 1)
    for sel in ATS_SUBMIT_SELECTORS:
        try:
            btn = page.locator(sel).last   # last = most likely the primary action
            if btn.is_visible(timeout=800):
                txt = (btn.inner_text(timeout=500) or btn.get_attribute("value") or "").lower()
                if any(w in txt for w in ["submit", "complete", "apply", "finish", "send"]):
                    btn.click()
                    _delay(1.5, 3)
                    return "submit"
        except Exception:
            continue
    for sel in ATS_NEXT_SELECTORS:
        try:
            btn = page.locator(sel).first
            if btn.is_visible(timeout=800):
                btn.click(); _delay(1, 2); return "next"
        except Exception:
            continue
    return "stuck"

def external_ats_apply(page, job: dict, pkg: dict, cv_path: str) -> bool:
    """
    Handle external ATS application.
    - Opens the apply link from the LinkedIn job page
    - If login page appears → pauses and asks you to log in
    - Auto-fills every form field it can
    - Walks through multi-step forms (Next → Next → ... → Submit)
    - Before final Submit, pauses for your review
    Returns True if submitted.
    """
    ats_name = _detect_ats(job.get("url", ""))
    print(f"      ATS detected: {ats_name}")
    print(f"      Opening LinkedIn job page to click Apply...")

    # Go back to the LinkedIn job page and click the Apply button
    page.goto(job["url"], wait_until="domcontentloaded", timeout=20000)
    _delay(1.5, 2.5)

    apply_page = None
    try:
        # LinkedIn Apply button usually opens a new tab
        with page.context.expect_page(timeout=8000) as p_info:
            apply_btn = page.locator(
                "button.jobs-apply-button:not(:has-text('Easy Apply')), "
                "a.jobs-apply-button"
            ).first
            apply_btn.click()
        apply_page = p_info.value
        apply_page.wait_for_load_state("domcontentloaded", timeout=15000)
        _delay(1.5, 2.5)
    except Exception:
        # Some jobs redirect in the same tab
        apply_page = page
        try:
            apply_btn = page.locator("button.jobs-apply-button, a.jobs-apply-button").first
            apply_btn.click()
            _delay(2, 3)
        except Exception:
            print(f"      ✗ Could not find Apply button")
            return False

    target = apply_page
    print(f"      ATS URL: {target.url[:80]}")

    # ── Pause for login if needed ──────────────────────────────────────────
    if _needs_login(target):
        print(f"\n{'─'*65}")
        print(f"  ACTION: Log in to {ats_name} in the browser window.")
        print(f"  Once you're on the application form, press Enter to continue.")
        print(f"{'─'*65}")
        input("  >> ")
        _delay(1, 2)

    # ── Multi-step form walk-through ───────────────────────────────────────
    print(f"      Filling form...")
    submitted = False
    for step in range(15):
        _delay(0.8, 1.5)

        # Check if we've landed on a confirmation page
        body = ""
        try:
            body = target.inner_text("body", timeout=3000).lower()
        except Exception:
            pass
        if any(w in body for w in ["application submitted", "application received",
                                    "thank you for applying", "successfully submitted",
                                    "we've received your application"]):
            submitted = True
            break

        # Fill all visible fields on this step
        _fill_ats_form(target, pkg, cv_path)

        # Try to advance
        result = _click_ats_next(target)
        if result == "submit":
            # Before final submit — pause for review
            print(f"\n{'─'*65}")
            print(f"  REVIEW: Form filled for {job['company']} — {job['title']}")
            print(f"  Check the browser. Fix anything wrong, then press Enter to submit.")
            print(f"  (Type 'skip' + Enter to skip this job without submitting.)")
            print(f"{'─'*65}")
            user_input = input("  >> ").strip().lower()
            if user_input == "skip":
                print(f"      Skipped by user.")
                return False
            # Click the submit button
            for sel in ATS_SUBMIT_SELECTORS:
                try:
                    btn = target.locator(sel).last
                    if btn.is_visible(timeout=1000):
                        txt = (btn.inner_text(timeout=500) or btn.get_attribute("value") or "").lower()
                        if any(w in txt for w in ["submit", "complete", "apply", "send"]):
                            btn.click(); _delay(2, 4); submitted = True; break
                except Exception:
                    continue
            break
        elif result == "stuck":
            print(f"      Stuck on step {step+1} — pausing for manual help")
            print(f"      Press Enter when you've moved past this step (or type 'skip'):")
            user_input = input("  >> ").strip().lower()
            if user_input == "skip":
                return False

    return submitted

# ── CV building ───────────────────────────────────────────────────────────────
def build_cv(pkg: dict, job: dict) -> str:
    """Generate tailored CV DOCX and return PDF path (falls back to DOCX)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        cv_text = read_cv()
        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Pt(36)
            section.left_margin = section.right_margin = Pt(54)

        def h(text, size=12, color=(0,51,102)):
            p = doc.add_paragraph()
            r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
            r.font.color.rgb = RGBColor(*color); return p

        def b(text, size=10.5):
            p = doc.add_paragraph()
            r = p.add_run(text); r.font.size = Pt(size); return p

        def rule():
            p = doc.add_paragraph("─" * 80)
            p.runs[0].font.size = Pt(7)
            p.runs[0].font.color.rgb = RGBColor(180,180,180)

        h("Faizan Fayyaz", size=18)
        b("London, UK  |  fznfayyaz@gmail.com  |  linkedin.com/in/faizanfayyaz")
        rule()

        h("Professional Summary")
        b(pkg.get("tailored_summary", ""))

        achs = pkg.get("key_achievements", [])
        if achs:
            rule(); h("Key Achievements")
            for a in achs:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(a).strip()).font.size = Pt(10.5)

        skills = pkg.get("skills_to_highlight", [])
        if skills:
            rule(); h("Core Skills"); b("  ·  ".join(skills))

        rule(); h("Career History")
        in_body = False
        for line in cv_text.split("\n"):
            if not in_body and re.search(r"\b(20\d{2}|19\d{2})\b", line):
                in_body = True
            if in_body and line.strip():
                b(line)

        cover = pkg.get("cover_letter", "")
        if cover:
            doc.add_page_break(); h(f"Cover Letter — {job['title']} at {job['company']}")
            rule()
            for para in cover.split("\n\n"):
                if para.strip(): b(para.strip())

        safe = lambda s: re.sub(r"[^\w]", "_", s)[:25]
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        docx_path = RESULTS_DIR / f"CV_{safe(job['company'])}_{safe(job['title'])}_{ts}.docx"
        doc.save(str(docx_path))

        pdf_path = docx_path.with_suffix(".pdf")
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(pdf_path))
            return str(pdf_path)
        except Exception:
            return str(docx_path)
    except Exception as e:
        print(f"      CV build error: {e} — using original CV")
        return CV_PDF_PATH

# ── Log result ────────────────────────────────────────────────────────────────
def log(job: dict, status: str, cv_path: str = ""):
    log_path = "applications.csv"
    exists = Path(log_path).exists()
    fields = ["timestamp","company","title","url","score","role_alignment","tech_alignment",
              "sponsorship","salary","company_quality","reason","status","cv_path"]
    with open(log_path, "a", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fields)
        if not exists: w.writeheader()
        w.writerow({
            "timestamp":       datetime.datetime.utcnow().isoformat(),
            "company":         job.get("company",""),
            "title":           job.get("title",""),
            "url":             job.get("url",""),
            "score":           job.get("score",""),
            "role_alignment":  job.get("role_alignment",""),
            "tech_alignment":  job.get("technology_alignment",""),
            "sponsorship":     job.get("sponsorship_probability",""),
            "salary":          job.get("salary_probability",""),
            "company_quality": job.get("company_quality",""),
            "reason":          job.get("reason",""),
            "status":          status,
            "cv_path":         cv_path,
        })

# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--top",        type=int, default=20,   help="Apply to top N jobs (default 20)")
    parser.add_argument("--min-score",  type=int, default=60,   help="Minimum score to apply (default 60)")
    parser.add_argument("--score-only", action="store_true",    help="Score and rank only, don't apply")
    parser.add_argument("--limit",      type=int, default=20,   help="Max jobs per search term (default 20)")
    parser.add_argument("--searches",   type=int, default=20,   help="How many search queries to generate (default 20)")
    args = parser.parse_args()

    if not CLAUDE_API_KEY:
        sys.exit("Set CLAUDE_API_KEY in .env")

    cv_text = read_cv()

    # ── Phase 0: Claude generates search queries from CV + MasterPrompt ───
    print(f"\n{'='*70}")
    print(f"  UK Job Hunt — autonomous mode")
    print(f"  top {args.top} · min score {args.min_score} · {'SCORE ONLY' if args.score_only else 'APPLY'}")
    print(f"{'='*70}")
    print(f"\n[Phase 0] Generating search queries from your CV + MasterPrompt...")

    searches = generate_searches(cv_text, n=args.searches)
    if searches:
        print(f"  Claude generated {len(searches)} search queries:")
        for i, s in enumerate(searches, 1):
            print(f"    {i:>2}. {s}")
    else:
        print(f"  Search generation failed — using fallback list")
        searches = FALLBACK_SEARCHES

    with sync_playwright() as pw:
        browser = pw.chromium.launch_persistent_context(
            user_data_dir=PROFILE_DIR, headless=False,
            args=["--disable-blink-features=AutomationControlled"],
            viewport={"width": 1280, "height": 900},
        )
        page = browser.new_page()

        # ── Phase 1: Collect job URLs ──────────────────────────────────────
        print(f"\n[Phase 1] Collecting job listings from LinkedIn...")
        all_urls = set()
        for i, kw in enumerate(searches, 1):
            url = build_search_url(kw)
            print(f"  [{i:>2}/{len(searches)}] '{kw}'", end="", flush=True)
            links = collect_job_links(page, url, limit=args.limit)
            new = [l for l in links if l not in all_urls]
            all_urls.update(new)
            print(f" → {len(new)} new  (total {len(all_urls)})")
            _delay(0.5, 1.0)

        all_urls = list(all_urls)
        print(f"\n  Total unique jobs found: {len(all_urls)}")

        # ── Phase 2: Scrape descriptions (with title pre-filter) ───────────
        print(f"\n[Phase 2] Scraping job descriptions...")
        jobs = []
        skipped_titles = 0
        for i, url in enumerate(all_urls, 1):
            job_id = url.split("/")[-1]
            print(f"  [{i:>3}/{len(all_urls)}]", end=" ", flush=True)
            job = scrape_job_page(page, url)
            if not job:
                print(f"{job_id} ✗ (failed to load)")
                continue
            # Pre-filter by title before wasting Claude tokens
            if not _title_passes(job["title"]):
                print(f"✗ BLOCKED TITLE: {job['title'][:50]}")
                skipped_titles += 1
                continue
            if len(job["description"]) < 100:
                print(f"{job_id} ✗ (no description)")
                continue
            jobs.append(job)
            print(f"✓ {job['company'][:30]} — {job['title'][:40]}")

        print(f"\n  Scraped {len(jobs)} jobs  ({skipped_titles} pre-filtered by title)")

        # ── Phase 3: Score with Claude (parallel, 10 workers) ─────────────
        print(f"\n[Phase 3] Scoring {len(jobs)} jobs with Claude (parallel)...")
        scored = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as ex:
            futures = {ex.submit(score_job, job.copy(), cv_text): job for job in jobs}
            done = 0
            for future in concurrent.futures.as_completed(futures):
                done += 1
                result = future.result()
                scored.append(result)
                bar = "█" * int(done / len(jobs) * 40)
                print(f"\r  [{bar:<40}] {done}/{len(jobs)}", end="", flush=True)
        print()

        # ── Phase 4: Rank ──────────────────────────────────────────────────
        apply_list = sorted(
            [j for j in scored if j["decision"] == "APPLY" and j["score"] >= args.min_score],
            key=lambda j: j["score"], reverse=True
        )[:args.top]

        skipped = [j for j in scored if j["decision"] != "APPLY" or j["score"] < args.min_score]

        # Save full results to CSV
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        results_csv = RESULTS_DIR / f"ranked_{ts}.csv"
        with open(results_csv, "w", newline="", encoding="utf-8") as f:
            fields = ["rank","score","decision","company","title","role_alignment",
                      "technology_alignment","sponsorship_probability","salary_probability",
                      "company_quality","reason","url"]
            w = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
            w.writeheader()
            for rank, j in enumerate(sorted(scored, key=lambda x: x["score"], reverse=True), 1):
                j["rank"] = rank; w.writerow(j)

        # Print ranked table
        print(f"\n{'─'*80}")
        print(f"  WILL APPLY: {len(apply_list)} jobs  (score ≥ {args.min_score}, out of {len(scored)} scored)")
        print(f"{'─'*80}")
        if apply_list:
            print(f"  {'#':>2}  {'Score':>5}  {'Role%':>5}  {'Tech%':>5}  {'Spns%':>5}  {'Company':<25} Title")
            print(f"  {'─'*2}  {'─'*5}  {'─'*5}  {'─'*5}  {'─'*5}  {'─'*25} {'─'*30}")
            for i, j in enumerate(apply_list, 1):
                print(f"  {i:>2}  {j['score']:>5}  "
                      f"{j.get('role_alignment',0):>5}  "
                      f"{j.get('technology_alignment',0):>5}  "
                      f"{j.get('sponsorship_probability',0):>5}  "
                      f"{j['company'][:25]:<25} {j['title'][:40]}")
                print(f"       → {j.get('reason','')[:78]}")
        else:
            print(f"  No jobs met the threshold of {args.min_score}.")
            print(f"  Try: python hunt.py --score-only --min-score 50")

        if skipped:
            print(f"\n  SKIPPED ({len(skipped)} jobs):")
            print(f"  {'Score':>5}  {'Company':<28} {'Title':<35} Reason")
            print(f"  {'─'*5}  {'─'*28} {'─'*35} {'─'*30}")
            for j in sorted(skipped, key=lambda x: x["score"], reverse=True):
                print(f"  {j['score']:>5}  {j['company'][:28]:<28} {j['title'][:35]:<35} {j.get('reason','')[:50]}")

        print(f"\n  Full results saved: {results_csv}")

        if args.score_only or not apply_list:
            if not apply_list:
                print(f"\n  No jobs met the threshold. Try --min-score 55 to lower the bar.")
            browser.close()
            return

        # ── Phase 5: Apply to top N ────────────────────────────────────────
        print(f"\n[Phase 5] Applying to {len(apply_list)} jobs automatically...\n")
        applied = skipped_apply = failed = 0

        for i, job in enumerate(apply_list, 1):
            print(f"\n  [{i}/{len(apply_list)}] {job['company']} — {job['title']}")
            print(f"      Score: {job['score']}  |  Role: {job.get('role_alignment',0)}%  "
                  f"Tech: {job.get('technology_alignment',0)}%  "
                  f"Sponsorship: {job.get('sponsorship_probability',0)}%")
            print(f"      {job['url']}")

            # Build tailored CV and application package (same for both ATS types)
            print(f"      Building tailored CV + cover letter with Claude...")
            pkg = build_application_package(job, cv_text)
            if not pkg:
                print(f"      ✗ Package generation failed")
                log(job, "FAILED")
                failed += 1
                continue

            cv_path = build_cv(pkg, job)

            if job["is_easy_apply"]:
                # ── LinkedIn Easy Apply: fully automatic ──────────────────
                print(f"      Type: LinkedIn Easy Apply (fully automatic)")
                page.goto(job["url"], wait_until="domcontentloaded", timeout=20000)
                _delay(1.5, 2)
                success = easy_apply(page, job, pkg, cv_path)
            else:
                # ── External ATS: auto-fill + pause for login/review ──────
                ats = _detect_ats(job.get("url", ""))
                print(f"      Type: External ATS ({ats})")
                print(f"      Bot will fill every field. You handle login + final review.")
                success = external_ats_apply(page, job, pkg, cv_path)

            if success:
                print(f"      ✓ SUBMITTED")
                log(job, "SUBMITTED", cv_path)
                applied += 1
            else:
                print(f"      ✗ FAILED / SKIPPED")
                log(job, "FAILED", cv_path)
                failed += 1

            _long_delay(2, 4)  # polite delay between applications

        print(f"\n{'='*65}")
        print(f"  HUNT COMPLETE")
        print(f"  Submitted: {applied}  |  Failed: {failed}  |  Skipped (external): {skipped_apply}")
        print(f"  Log: applications.csv")
        print(f"{'='*65}")

        input("\nPress Enter to close browser...")
        browser.close()

if __name__ == "__main__":
    main()
